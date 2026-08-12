import json
from datetime import datetime, timedelta

from django.contrib.auth import login as auth_login, logout as auth_logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse, HttpResponseForbidden, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.views.decorators.http import require_POST, require_http_methods

from . import seo_engine
from . import llm_engine
from .models import AuditRecord, LearnerProfile, QuizAttempt, QuizQuestion, ChatThread, ChatMessage


# =============================================================================
# QUIZ CATALOG — static content, but now data-driven so it can be filtered
# server-side by each learner's chosen difficulty level and interests.
# Balanced: 4 easy / 4 medium / 4 hard, spread across all 5 topic categories.
# =============================================================================
QUIZ_CATALOG = [
    # --- Easy ---
    {'id': 'perf-basics', 'title': 'Performance Basics',
     'desc': 'Core Web Vitals, load time, and rendering metrics',
     'icon': 'fa-gauge-high', 'accent': '#5b8def', 'difficulty': 'easy',
     'category': 'performance', 'questions': 15, 'minutes': 10, 'progress': 0},
    {'id': 'testing-concepts', 'title': 'Testing Concepts',
     'desc': 'Test types, test cases, boundary analysis, and strategies',
     'icon': 'fa-vial', 'accent': '#00e0a1', 'difficulty': 'easy',
     'category': 'testing', 'questions': 12, 'minutes': 8, 'progress': 0},
    {'id': 'seo-keyword-basics', 'title': 'Keyword & Content Basics',
     'desc': 'Keyword research, on-page content, and search intent fundamentals',
     'icon': 'fa-font', 'accent': '#ff8fa3', 'difficulty': 'easy',
     'category': 'seo', 'questions': 14, 'minutes': 9, 'progress': 0},
    {'id': 'responsive-web-design', 'title': 'Responsive Web Design',
     'desc': 'Media queries, flexible layouts, and mobile-first fundamentals',
     'icon': 'fa-mobile-screen', 'accent': '#c084fc', 'difficulty': 'easy',
     'category': 'web', 'questions': 13, 'minutes': 9, 'progress': 0},

    # --- Medium ---
    {'id': 'seo-fundamentals', 'title': 'SEO Fundamentals',
     'desc': 'On-page SEO, meta tags, heading structure, and indexing',
     'icon': 'fa-magnifying-glass', 'accent': '#ff5c6c', 'difficulty': 'medium',
     'category': 'seo', 'questions': 20, 'minutes': 15, 'progress': 0},
    {'id': 'automation-tools', 'title': 'Automation Testing',
     'desc': 'Test automation tools, CI/CD, and automated analysis pipelines',
     'icon': 'fa-robot', 'accent': '#38bdf8', 'difficulty': 'medium',
     'category': 'testing', 'questions': 10, 'minutes': 7, 'progress': 0},
    {'id': 'core-web-vitals-deep-dive', 'title': 'Core Web Vitals Deep Dive',
     'desc': 'LCP, INP, and CLS: measuring and improving real user experience',
     'icon': 'fa-chart-line', 'accent': '#4dd4ac', 'difficulty': 'medium',
     'category': 'performance', 'questions': 18, 'minutes': 14, 'progress': 0},
    {'id': 'usability-reliability', 'title': 'Usability & Reliability',
     'desc': 'Evaluating ease of use, error tolerance, and system dependability',
     'icon': 'fa-user-check', 'accent': '#f4a742', 'difficulty': 'medium',
     'category': 'nfr', 'questions': 16, 'minutes': 12, 'progress': 0},

    # --- Hard ---
    {'id': 'nfr-mastery', 'title': 'NFR Mastery',
     'desc': 'Non-functional requirements: reliability, usability, scalability',
     'icon': 'fa-shield-halved', 'accent': '#ffaa3b', 'difficulty': 'hard',
     'category': 'nfr', 'questions': 25, 'minutes': 20, 'progress': 0},
    {'id': 'web-accessibility', 'title': 'Web Accessibility',
     'desc': 'WCAG guidelines, ARIA labels, screen readers, and compliance',
     'icon': 'fa-universal-access', 'accent': '#a78bfa', 'difficulty': 'hard',
     'category': 'web', 'questions': 18, 'minutes': 12, 'progress': 0},
    {'id': 'technical-seo-audits', 'title': 'Technical SEO Audits',
     'desc': 'Crawl budget, canonical tags, structured data, and indexability audits',
     'icon': 'fa-server', 'accent': '#e05656', 'difficulty': 'hard',
     'category': 'seo', 'questions': 22, 'minutes': 18, 'progress': 0},
    {'id': 'load-stress-testing', 'title': 'Load & Stress Testing',
     'desc': 'Simulating traffic spikes, breakpoints, and performance under load',
     'icon': 'fa-weight-hanging', 'accent': '#2f9e94', 'difficulty': 'hard',
     'category': 'performance', 'questions': 20, 'minutes': 18, 'progress': 0},

    # --- Short "quick check" quizzes — ensure there's always something
    # available even when a learner has very little daily-goal time left.
    {'id': 'quick-meta-tags', 'title': 'Quick Meta Tags Check',
     'desc': 'A fast run-through of title tags, meta descriptions, and canonical basics',
     'icon': 'fa-tags', 'accent': '#ff8fa3', 'difficulty': 'easy',
     'category': 'seo', 'questions': 8, 'minutes': 5, 'progress': 0},
    {'id': 'smoke-testing-basics', 'title': 'Smoke Testing Basics',
     'desc': 'What smoke tests check, when to run them, and why they matter',
     'icon': 'fa-fire', 'accent': '#00e0a1', 'difficulty': 'easy',
     'category': 'testing', 'questions': 9, 'minutes': 6, 'progress': 0},
    {'id': 'mobile-viewport-essentials', 'title': 'Mobile Viewport Essentials',
     'desc': 'Viewport meta tags, breakpoints, and mobile-friendliness checks',
     'icon': 'fa-mobile-screen-button', 'accent': '#c084fc', 'difficulty': 'medium',
     'category': 'web', 'questions': 11, 'minutes': 9, 'progress': 0},
    {'id': 'canonical-redirects-quickfire', 'title': 'Canonical & Redirects Quickfire',
     'desc': 'A rapid-fire round on canonical tags, 301s, and redirect chains',
     'icon': 'fa-arrows-turn-right', 'accent': '#e05656', 'difficulty': 'medium',
     'category': 'seo', 'questions': 10, 'minutes': 8, 'progress': 0},
]

# Maps free-text interests (whatever a learner types on their Profile page,
# e.g. "Load Testing", "Core Web Vitals") to the fixed quiz categories above,
# so the Quiz page can actually be filtered by what someone said they care
# about — not just decorative tags.
INTEREST_CATEGORY_KEYWORDS = {
    'performance': ['performance', 'speed', 'load time', 'load testing', 'core web vital',
                    'vitals', 'latency', 'rendering'],
    'seo': ['seo', 'search engine', 'meta tag', 'metadata', 'keyword', 'ranking', 'indexing'],
    'testing': ['test', 'testing', 'qa', 'quality assurance', 'automation', 'boundary',
                'test case', 'ci/cd', 'pipeline'],
    'nfr': ['nfr', 'non-functional', 'nonfunctional', 'reliability', 'scalability',
            'usability', 'availability'],
    'web': ['web standard', 'accessibility', 'wcag', 'aria', 'screen reader', 'compliance'],
}


CATEGORY_LABELS = {
    'performance': 'Performance',
    'seo': 'SEO',
    'testing': 'Testing',
    'nfr': 'NFR',
    'web': 'Web Standards',
}


def match_categories_from_interests(interests):
    """Best-effort keyword match — returns the set of quiz categories that
    at least one of the learner's interests seems to be about."""
    matched = set()
    for interest in interests or []:
        text = interest.lower()
        for category, keywords in INTEREST_CATEGORY_KEYWORDS.items():
            if any(kw in text for kw in keywords):
                matched.add(category)
    return matched


# =============================================================================
# PUBLIC PAGES
# =============================================================================
def home(request):
    all_records = AuditRecord.objects.all()
    site_stats = {
        'websites_tested': all_records.count(),
        'missions_completed': sum(len(r.completed_indexes or []) for r in all_records),
        'registered_learners': User.objects.count(),
    }
    return render(request, 'home.html', {'site_stats': site_stats})


# =============================================================================
# AUTH: signup / login / logout
# =============================================================================
def signup(request):
    if request.user.is_authenticated:
        return redirect('dashboard')

    if request.method == 'POST':
        username = (request.POST.get('username') or '').strip()
        email = (request.POST.get('email') or '').strip()
        password1 = request.POST.get('password1') or ''
        password2 = request.POST.get('password2') or ''

        form = UserCreationForm(data={
            'username': username, 'password1': password1, 'password2': password2,
        })
        if form.is_valid():
            user = form.save(commit=False)
            user.email = email
            user.save()
            auth_login(request, user)
            messages.success(request, f'Welcome, {user.username}! Your account is ready.')
            return redirect('dashboard')
        else:
            for field_errors in form.errors.values():
                for err in field_errors:
                    messages.error(request, err)

    return render(request, 'signup.html')


def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')

    if request.method == 'POST':
        username = (request.POST.get('username') or '').strip()
        password = request.POST.get('password') or ''
        user = authenticate(request, username=username, password=password)
        if user is not None:
            auth_login(request, user)
            next_url = request.POST.get('next') or request.GET.get('next') or 'dashboard'
            return redirect(next_url)
        messages.error(request, 'Incorrect username or password.')

    return render(request, 'login.html', {'next': request.GET.get('next', '')})


@require_http_methods(["GET", "POST"])
def logout_view(request):
    auth_logout(request)
    return redirect('home')


# =============================================================================
# PROTECTED PAGES (require login)
# =============================================================================
@login_required
def chatbot(request):
    threads = ChatThread.objects.filter(user=request.user)
    active_thread = threads.first()  # most recently updated, or None
    active_messages = active_thread.messages.all() if active_thread else []
    context = {
        'threads': threads,
        'active_thread': active_thread,
        'active_messages': active_messages,
        'llm_configured': llm_engine.is_configured(),
    }
    return render(request, 'chatbot.html', context)


@login_required
@require_POST
def chat_new_thread(request):
    thread = ChatThread.objects.create(user=request.user)
    return JsonResponse({'ok': True, 'thread_id': thread.id, 'title': thread.title})


@login_required
def chat_thread_messages(request, thread_id):
    thread = get_object_or_404(ChatThread, pk=thread_id, user=request.user)
    messages_data = [
        {'role': m.role, 'content': m.content, 'time': m.created_at.strftime('%I:%M %p')}
        for m in thread.messages.all()
    ]
    return JsonResponse({'ok': True, 'thread_id': thread.id, 'title': thread.title, 'messages': messages_data})


@login_required
@require_POST
def chat_send_message(request, thread_id):
    thread = get_object_or_404(ChatThread, pk=thread_id, user=request.user)
    try:
        body = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'ok': False, 'error': 'Invalid request.'}, status=400)

    content = (body.get('content') or '').strip()
    if not content:
        return JsonResponse({'ok': False, 'error': 'Message cannot be empty.'}, status=400)
    if len(content) > 4000:
        return JsonResponse({'ok': False, 'error': 'Message is too long (max 4000 characters).'}, status=400)

    is_first_message = not thread.messages.exists()
    ChatMessage.objects.create(thread=thread, role='user', content=content)

    if is_first_message:
        thread.title = llm_engine.make_thread_title(content)
        thread.save(update_fields=['title'])

    # Bound the context sent to the API — last 20 messages is plenty for a
    # tutoring chat and keeps token usage (and cost) predictable.
    recent = list(thread.messages.order_by('-created_at')[:20])[::-1]
    history = [{'role': m.role, 'content': m.content} for m in recent]

    reply_text, ok = llm_engine.ask_tutor(history)
    ChatMessage.objects.create(thread=thread, role='assistant', content=reply_text)
    thread.save()  # bump updated_at so this thread sorts to the top

    return JsonResponse({
        'ok': True,
        'llm_ok': ok,
        'reply': reply_text,
        'thread_id': thread.id,
        'thread_title': thread.title,
    })


@login_required
@require_POST
def chat_delete_thread(request, thread_id):
    thread = get_object_or_404(ChatThread, pk=thread_id, user=request.user)
    thread.delete()
    return JsonResponse({'ok': True})


@login_required
def chat_export_thread(request, thread_id):
    """Exports a chat thread as a real .docx file (not client-side .txt),
    so the download opens directly in Word with proper formatting."""
    thread = get_object_or_404(ChatThread, pk=thread_id, user=request.user)

    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(thread.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"EduTest AI · LLM Tutor conversation\n"
        f"Exported {timezone.now().strftime('%d %b %Y, %I:%M %p')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_paragraph()  # spacer

    for msg in thread.messages.all():
        speaker = "You" if msg.role == "user" else "LLM Tutor"
        p = doc.add_paragraph()
        speaker_run = p.add_run(f"{speaker} · {msg.created_at.strftime('%I:%M %p')}")
        speaker_run.bold = True
        speaker_run.font.size = Pt(10)
        speaker_run.font.color.rgb = (
            RGBColor(0x2F, 0x6F, 0xED) if msg.role == "user" else RGBColor(0x1A, 0xA1, 0x7A)
        )
        body = doc.add_paragraph(msg.content)
        body.paragraph_format.space_after = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = "".join(c for c in thread.title if c.isalnum() or c in " -_").strip() or "chat-export"
    filename = f"{safe_title[:60]}.docx"

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _score_trend_buckets(records):
    """Buckets a user's real AuditRecord history into week/month/year views
    for the 'Recent URL Test Trend' chart. Only buckets that actually
    contain an audit are included — no fabricated zero-score days."""
    now = timezone.now()

    # WEEK: average score per day, last 7 calendar days
    week_buckets = {}
    for r in records:
        day = r.created_at.date()
        if (now.date() - day).days < 7:
            week_buckets.setdefault(day, []).append(r.overall)
    week_labels, week_data = [], []
    for i in range(6, -1, -1):
        day = (now - timedelta(days=i)).date()
        if day in week_buckets:
            week_labels.append(day.strftime('%a %d'))
            week_data.append(round(sum(week_buckets[day]) / len(week_buckets[day])))

    # MONTH: average score per week, last 4 weeks
    month_buckets = {}
    for r in records:
        delta_days = (now.date() - r.created_at.date()).days
        if 0 <= delta_days < 28:
            month_buckets.setdefault(delta_days // 7, []).append(r.overall)
    month_labels, month_data = [], []
    for week_idx in range(3, -1, -1):
        if week_idx in month_buckets:
            label = 'This Week' if week_idx == 0 else f'{week_idx} Week{"s" if week_idx > 1 else ""} Ago'
            month_labels.append(label)
            month_data.append(round(sum(month_buckets[week_idx]) / len(month_buckets[week_idx])))

    # YEAR: average score per month, last 12 months
    year_buckets = {}
    for r in records:
        key = (r.created_at.year, r.created_at.month)
        year_buckets.setdefault(key, []).append(r.overall)
    months_back = []
    cursor_year, cursor_month = now.year, now.month
    for _ in range(12):
        months_back.append((cursor_year, cursor_month))
        cursor_month -= 1
        if cursor_month == 0:
            cursor_month = 12
            cursor_year -= 1
    year_labels, year_data = [], []
    for (y, m) in reversed(months_back):
        if (y, m) in year_buckets:
            year_labels.append(datetime(y, m, 1).strftime('%b'))
            year_data.append(round(sum(year_buckets[(y, m)]) / len(year_buckets[(y, m)])))

    return {
        'week': {'labels': week_labels, 'data': week_data},
        'month': {'labels': month_labels, 'data': month_data},
        'year': {'labels': year_labels, 'data': year_data},
    }


def _user_activity_stats(user):
    """Every real, per-user number shown on the Dashboard and Profile pages
    is computed here from actual AuditRecord rows — nothing hardcoded."""
    records = list(AuditRecord.objects.filter(user=user))  # already ordered newest-first
    total_audits = len(records)
    total_xp = sum(r.total_xp for r in records)
    total_missions = sum(r.completed_count for r in records)
    total_possible_missions = sum(r.total_count for r in records)
    avg_score = round(sum(r.overall for r in records) / total_audits) if total_audits else 0

    earned_badge_ids = set()
    for r in records:
        earned_badge_ids.update(r.earned_badges or [])
    badges_with_status = [
        dict(b, earned=b['id'] in earned_badge_ids) for b in seo_engine.BADGES
    ]

    # Oldest -> newest, last 12 audits, for the score-trend line chart
    trend_records = sorted(records, key=lambda r: r.created_at)[-12:]
    score_trend = {
        'labels': [r.created_at.strftime('%b %d') for r in trend_records],
        'data': [r.overall for r in trend_records],
    }

    # Average of each SEO category across every audit this user has run,
    # for the category-performance chart.
    category_totals, category_counts = {}, {}
    for r in records:
        for cat, val in (r.categories or {}).items():
            category_totals[cat] = category_totals.get(cat, 0) + val
            category_counts[cat] = category_counts.get(cat, 0) + 1
    category_averages = {
        cat: round(category_totals[cat] / category_counts[cat]) for cat in category_totals
    }

    return {
        'records': records,
        'total_audits': total_audits,
        'total_xp': total_xp,
        'avg_score': avg_score,
        'total_missions': total_missions,
        'total_possible_missions': total_possible_missions,
        'badges_with_status': badges_with_status,
        'badges_earned_count': len(earned_badge_ids),
        'score_trend': score_trend,
        'category_averages': category_averages,
        'audit_dates': [r.created_at.date().isoformat() for r in records],
    }


@login_required
def dashboard(request):
    stats = _user_activity_stats(request.user)
    context = {
        'recent_records': stats['records'][:8],
        'total_audits': stats['total_audits'],
        'total_xp': stats['total_xp'],
        'avg_score': stats['avg_score'],
        'total_missions': stats['total_missions'],
        'total_possible_missions': stats['total_possible_missions'],
        'badges_with_status': stats['badges_with_status'],
        'badges_earned_count': stats['badges_earned_count'],
        'score_trend_json': stats['score_trend'],
        'category_averages_json': stats['category_averages'],
    }
    return render(request, 'dashboard.html', context)


def _daily_goal_status(user, learner_profile):
    """Real daily-goal progress computed from actual QuizAttempt rows.
    Auto-unlocks the goal (clears daily_goal_locked_at) the moment enough
    real minutes have been logged today — no manual reset needed."""
    today = timezone.localdate()
    minutes_today = sum(
        a.minutes for a in QuizAttempt.objects.filter(user=user, completed_at__date=today)
    )
    goal = learner_profile.daily_goal_minutes
    goal_reached = minutes_today >= goal
    is_locked = bool(learner_profile.daily_goal_locked_at) and not goal_reached

    if goal_reached and learner_profile.daily_goal_locked_at:
        learner_profile.daily_goal_locked_at = None
        learner_profile.save(update_fields=['daily_goal_locked_at'])

    return {
        'minutes_completed_today': minutes_today,
        'goal_minutes': goal,
        'remaining_minutes': max(0, goal - minutes_today),
        'goal_reached': goal_reached,
        'is_locked': is_locked,
    }


@login_required
def quiz(request):
    learner_profile, _ = LearnerProfile.objects.get_or_create(user=request.user)
    goal_status = _daily_goal_status(request.user, learner_profile)

    selected_quiz_levels = learner_profile.quiz_difficulties  # e.g. ['easy', 'hard']
    # No levels selected = nothing to filter by, so show everything rather
    # than an empty page. "?level=all" also forces showing everything.
    show_all_difficulty = request.GET.get('level') == 'all' or not selected_quiz_levels

    matched_interest_categories = match_categories_from_interests(learner_profile.interests)
    # Same fallback logic: no interests set, or none of them matched a real
    # topic, means there's nothing sensible to filter by — show everything.
    # "?interests=all" also forces showing everything regardless.
    show_all_interests = request.GET.get('interests') == 'all' or not matched_interest_categories

    # Time-budget filter: once today's goal is met, stop restricting by time
    # (there's nothing left to protect them from) — otherwise only show
    # quizzes that actually fit in the minutes they have left today.
    show_all_time = (
        request.GET.get('time') == 'all'
        or goal_status['goal_reached']
        or goal_status['remaining_minutes'] <= 0
    )

    cards = QUIZ_CATALOG
    if not show_all_difficulty:
        cards = [c for c in cards if c['difficulty'] in selected_quiz_levels]
    if not show_all_interests:
        cards = [c for c in cards if c['category'] in matched_interest_categories]
    if not show_all_time:
        cards = [c for c in cards if c['minutes'] <= goal_status['remaining_minutes']]

    context = {
        'quiz_cards': cards,
        'show_all': show_all_difficulty,
        'selected_quiz_levels': selected_quiz_levels,
        'difficulty_labels': learner_profile.difficulty_labels,
        'has_interests': bool(learner_profile.interests),
        'show_all_interests': show_all_interests,
        'matched_interest_categories': sorted(matched_interest_categories),
        'matched_interest_labels': [CATEGORY_LABELS.get(c, c) for c in sorted(matched_interest_categories)],
        'show_all_time': show_all_time,
        'goal_status': goal_status,
    }
    return render(request, 'quiz.html', context)


@login_required
@require_POST
def quiz_check_answer(request, question_id):
    """Immediate right/wrong feedback for one question — called the moment
    a learner picks an option, not at quiz submission time. Deliberately a
    separate endpoint from quiz_questions() so the correct answer for a
    question is only revealed after the learner has actually answered it,
    not bundled into the initial question payload."""
    question = get_object_or_404(QuizQuestion, pk=question_id)
    try:
        body = json.loads(request.body.decode('utf-8'))
        selected = int(body.get('selected'))
    except (json.JSONDecodeError, UnicodeDecodeError, TypeError, ValueError):
        return JsonResponse({'ok': False, 'error': 'Invalid request.'}, status=400)

    return JsonResponse({
        'ok': True,
        'correct': selected == question.correct_index,
        'correct_index': question.correct_index,
        'explanation': question.explanation,
    })


@login_required
def quiz_questions(request, quiz_id):
    """Real per-topic questions for one quiz, served from the database
    (QuizQuestion) — never the correct answer, only the choices, so the
    answer key can't be read out of the page source or network tab."""
    matched_quiz = next((c for c in QUIZ_CATALOG if c['id'] == quiz_id), None)
    if not matched_quiz:
        return JsonResponse({'ok': False, 'error': 'Unknown quiz.'}, status=404)

    questions = list(QuizQuestion.objects.filter(quiz_id=quiz_id).order_by('order'))
    if not questions:
        return JsonResponse({'ok': False, 'error': 'This quiz has no questions yet.'}, status=404)

    return JsonResponse({
        'ok': True,
        'quiz_id': quiz_id,
        'quiz_title': matched_quiz['title'],
        'questions': [
            {'id': q.id, 'order': q.order, 'text': q.text, 'choices': q.choices}
            for q in questions
        ],
    })


@login_required
@require_POST
def complete_quiz(request):
    """Called when a learner finishes a quiz (see quiz.js finishQuiz()).
    Grades the submitted answers server-side against QuizQuestion.correct_index
    (the client never receives the answer key), logs a real QuizAttempt using
    the quiz's own authoritative minutes value — never trusts a client-sent
    duration — then recomputes daily-goal status, auto-unlocking the goal
    slider if today's target is now met."""
    try:
        body = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'ok': False, 'error': 'Invalid request.'}, status=400)

    quiz_id = body.get('quiz_id')
    matched_quiz = next((c for c in QUIZ_CATALOG if c['id'] == quiz_id), None)
    if not matched_quiz:
        return JsonResponse({'ok': False, 'error': 'Unknown quiz.'}, status=400)

    # answers: {"<question_id>": <selected_choice_index>, ...}. Keys arrive
    # as strings over JSON; coerce to int below for the DB id lookup.
    answers = body.get('answers') or {}
    questions = list(QuizQuestion.objects.filter(quiz_id=quiz_id))

    correct_count = 0
    for q in questions:
        selected = answers.get(str(q.id))
        if selected is not None:
            try:
                if int(selected) == q.correct_index:
                    correct_count += 1
            except (TypeError, ValueError):
                pass

    total_questions = len(questions)
    score_percent = round((correct_count / total_questions) * 100) if total_questions else None

    QuizAttempt.objects.create(
        user=request.user,
        quiz_id=matched_quiz['id'],
        quiz_title=matched_quiz['title'],
        minutes=matched_quiz['minutes'],
    )

    learner_profile, _ = LearnerProfile.objects.get_or_create(user=request.user)
    goal_status = _daily_goal_status(request.user, learner_profile)

    return JsonResponse({
        'ok': True,
        'quiz_title': matched_quiz['title'],
        'minutes_logged': matched_quiz['minutes'],
        'goal_status': goal_status,
        'score_percent': score_percent,
        'correct_count': correct_count,
        'total_questions': total_questions,
    })


@login_required
def analysis(request):
    stats = _user_activity_stats(request.user)
    trend_buckets = _score_trend_buckets(stats['records'])
    context = {
        'has_audit_history': stats['total_audits'] > 0,
        'score_trend_buckets_json': trend_buckets,
        'category_averages_json': stats['category_averages'],
    }
    return render(request, 'analysis.html', context)


@login_required
def profile(request):
    learner_profile, _ = LearnerProfile.objects.get_or_create(user=request.user)

    if request.method == 'POST':
        request.user.first_name = (request.POST.get('first_name') or '').strip()
        request.user.last_name = (request.POST.get('last_name') or '').strip()
        request.user.email = (request.POST.get('email') or '').strip()
        request.user.save()
        messages.success(request, 'Profile updated.')
        return redirect('profile')

    stats = _user_activity_stats(request.user)
    goal_status = _daily_goal_status(request.user, learner_profile)
    context = {
        'total_audits': stats['total_audits'],
        'total_xp': stats['total_xp'],
        'avg_score': stats['avg_score'],
        'total_missions': stats['total_missions'],
        'badges_unlocked': stats['badges_earned_count'],
        'audit_dates_json': stats['audit_dates'],
        'learner_profile': learner_profile,
        'goal_status': goal_status,
        # These 5 exact names guarantee a reliable match against quiz
        # categories (see match_categories_from_interests) — picking one
        # from the dropdown always filters the Quiz page correctly, unlike
        # free-typed text which is only a best-effort keyword match.
        'interest_topic_choices': list(CATEGORY_LABELS.values()),
    }
    return render(request, 'profile.html', context)


@login_required
@require_POST
def update_preferences(request):
    """AJAX endpoint the Profile page's Learning Preferences panel calls
    the moment a toggle/slider changes — persists instantly, no save button."""
    try:
        body = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'ok': False, 'error': 'Invalid request.'}, status=400)

    learner_profile, _ = LearnerProfile.objects.get_or_create(user=request.user)
    valid_difficulties = dict(LearnerProfile.DIFFICULTY_CHOICES)
    valid_styles = dict(LearnerProfile.LEARNING_STYLE_CHOICES)

    if 'difficulty_levels' in body:
        vals = body['difficulty_levels']
        if not isinstance(vals, list) or not all(v in valid_difficulties for v in vals):
            return JsonResponse({'ok': False, 'error': 'Invalid difficulty levels.'}, status=400)
        # De-dupe but keep a stable order (Beginner, Intermediate, Advanced)
        ordered = [k for k in valid_difficulties if k in vals]
        learner_profile.difficulty_levels = ordered

    if 'learning_style' in body:
        val = body['learning_style']
        if val not in valid_styles:
            return JsonResponse({'ok': False, 'error': 'Invalid learning style.'}, status=400)
        learner_profile.learning_style = val

    if 'daily_goal_minutes' in body:
        goal_status = _daily_goal_status(request.user, learner_profile)
        if goal_status['is_locked']:
            return JsonResponse({
                'ok': False,
                'error': (f"Your daily goal is locked until you complete "
                          f"{goal_status['remaining_minutes']} more minute(s) of quizzes today."),
            }, status=400)
        try:
            val = int(body['daily_goal_minutes'])
        except (TypeError, ValueError):
            return JsonResponse({'ok': False, 'error': 'Invalid daily goal.'}, status=400)
        learner_profile.daily_goal_minutes = max(10, min(120, val))
        # Locking happens the instant a new goal is picked — it stays
        # locked until real QuizAttempt minutes today meet it.
        learner_profile.daily_goal_locked_at = timezone.now()

    if 'interests' in body:
        vals = body['interests']
        if not isinstance(vals, list) or not all(isinstance(v, str) for v in vals):
            return JsonResponse({'ok': False, 'error': 'Invalid interests.'}, status=400)
        # Trim, drop empties/dupes (case-insensitive), cap length so this
        # can't be abused to store arbitrarily large blobs.
        seen = set()
        cleaned = []
        for v in vals:
            v = v.strip()[:40]
            if v and v.lower() not in seen:
                seen.add(v.lower())
                cleaned.append(v)
        learner_profile.interests = cleaned[:20]

    learner_profile.save()
    return JsonResponse({
        'ok': True,
        'difficulty_levels': learner_profile.difficulty_levels,
        'difficulty_labels': learner_profile.difficulty_labels,
        'learning_style': learner_profile.learning_style,
        'daily_goal_minutes': learner_profile.daily_goal_minutes,
        'interests': learner_profile.interests,
        'goal_status': _daily_goal_status(request.user, learner_profile),
    })


# =============================================================================
# Website Analyzer: runs the real crawler + XGBoost model + issue detector,
# and saves the result as a persistent, per-user AuditRecord.
# =============================================================================
@login_required
@require_POST
def analyze_url(request):
    try:
        body = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'ok': False, 'error': 'Invalid request body.'}, status=400)

    url = (body.get('url') or '').strip()
    if not url:
        return JsonResponse({'ok': False, 'error': 'Please provide a URL.'}, status=400)
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url

    result = seo_engine.run_full_audit(url)
    status_code = 200 if result.get('ok') else 422

    if result.get('ok'):
        game = result['game']
        record = AuditRecord.objects.create(
            user=request.user,
            url=url,
            overall=result['overall'],
            ml_score=result['ml_score'],
            grade=result['grade'],
            categories=result['categories'],
            issues=game['tasks'],
            errors_count=result['errors_count'],
            warnings_count=result['warnings_count'],
            total_xp=0,
            max_xp=game['max_xp'],
            completed_indexes=[],
            earned_badges=[],
        )
        result['record_id'] = record.id
        result['console_url'] = f'/analysis/console/{record.id}/'

    return JsonResponse(result, status=status_code)


def _level_payload(total_xp):
    name, color, progress, next_at = seo_engine.get_level(total_xp)
    return {'name': name, 'color': color, 'progress': progress, 'next_at': next_at}


def _record_payload(record, read_only=False):
    return {
        'record_id': record.id,
        'url': record.url,
        'tasks': record.issues,
        'completed_indexes': record.completed_indexes,
        'total_xp': record.total_xp,
        'max_xp': record.max_xp,
        'completed_count': record.completed_count,
        'total_count': record.total_count,
        'level': _level_payload(record.total_xp),
        'badges': seo_engine.BADGES,
        'earned_badge_ids': record.earned_badges,
        'read_only': read_only,
    }


# =============================================================================
# Game Console — a dedicated full page per audit record. Regular users only
# ever see their own records; staff can additionally view (read-only) any
# student's console for support/grading purposes.
# =============================================================================
@login_required
def game_console(request, record_id):
    record = get_object_or_404(AuditRecord, pk=record_id)
    is_owner = record.user == request.user
    if not is_owner and not request.user.is_staff:
        return HttpResponseForbidden("You don't have permission to view this record.")
    context = {
        'has_game': True,
        'game_payload': _record_payload(record, read_only=not is_owner),
        'viewing_as_staff': not is_owner,
        'record_owner': record.user,
    }
    return render(request, 'game_console.html', context)


@login_required
def game_console_latest(request):
    """Convenience redirect: /analysis/console/ -> the user's most recent audit."""
    record = AuditRecord.objects.filter(user=request.user).first()
    if not record:
        return render(request, 'game_console.html', {'has_game': False})
    return redirect('game_console', record_id=record.id)


@login_required
@require_POST
def validate_task(request, record_id):
    record = get_object_or_404(AuditRecord, pk=record_id, user=request.user)
    tasks = record.issues

    try:
        body = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'ok': False, 'error': 'Invalid request body.'}, status=400)

    try:
        idx = int(body.get('task_index'))
    except (TypeError, ValueError):
        idx = -1
    if idx < 0 or idx >= len(tasks):
        return JsonResponse({'ok': False, 'error': 'Invalid task.'}, status=400)

    completed = set(record.completed_indexes)
    if idx in completed:
        return JsonResponse({'ok': False, 'error': 'Task already completed.'}, status=400)

    issue = tasks[idx]
    correct, message = seo_engine.validate_task_action(issue, body)
    if not correct:
        return JsonResponse({'ok': True, 'correct': False, 'message': message})

    xp_val = seo_engine.xp_for_task(issue['imp'])
    completed.add(idx)

    completed_keys = {f'task_{i}' for i in completed}
    prev_badges = set(record.earned_badges)
    earned_badges = seo_engine.compute_earned_badges(tasks, completed_keys)
    newly_earned_ids = earned_badges - prev_badges

    record.completed_indexes = sorted(completed)
    record.total_xp = record.total_xp + xp_val
    record.earned_badges = sorted(earned_badges)
    record.save(update_fields=['completed_indexes', 'total_xp', 'earned_badges', 'updated_at'])

    return JsonResponse({
        'ok': True,
        'correct': True,
        'message': message,
        'task_index': idx,
        'xp_awarded': xp_val,
        'total_xp': record.total_xp,
        'completed_count': record.completed_count,
        'total_count': record.total_count,
        'level': _level_payload(record.total_xp),
        'newly_earned_badges': [b for b in seo_engine.BADGES if b['id'] in newly_earned_ids],
        'earned_badge_ids': record.earned_badges,
    })


@login_required
@require_POST
def undo_task(request, record_id):
    record = get_object_or_404(AuditRecord, pk=record_id, user=request.user)
    tasks = record.issues

    try:
        body = json.loads(request.body.decode('utf-8'))
        idx = int(body.get('task_index'))
    except (json.JSONDecodeError, UnicodeDecodeError, TypeError, ValueError):
        return JsonResponse({'ok': False, 'error': 'Invalid request.'}, status=400)

    completed = set(record.completed_indexes)
    if idx not in completed or idx >= len(tasks):
        return JsonResponse({'ok': False, 'error': 'Task was not completed.'}, status=400)

    xp_val = seo_engine.xp_for_task(tasks[idx]['imp'])
    completed.discard(idx)
    completed_keys = {f'task_{i}' for i in completed}
    earned_badges = seo_engine.compute_earned_badges(tasks, completed_keys)

    record.completed_indexes = sorted(completed)
    record.total_xp = max(0, record.total_xp - xp_val)
    record.earned_badges = sorted(earned_badges)
    record.save(update_fields=['completed_indexes', 'total_xp', 'earned_badges', 'updated_at'])

    return JsonResponse({
        'ok': True,
        'task_index': idx,
        'total_xp': record.total_xp,
        'completed_count': record.completed_count,
        'total_count': record.total_count,
        'level': _level_payload(record.total_xp),
        'earned_badge_ids': record.earned_badges,
    })


@login_required
@require_POST
def reset_game(request, record_id):
    record = get_object_or_404(AuditRecord, pk=record_id, user=request.user)

    record.completed_indexes = []
    record.total_xp = 0
    record.earned_badges = []
    record.save(update_fields=['completed_indexes', 'total_xp', 'earned_badges', 'updated_at'])

    return JsonResponse({
        'ok': True,
        'total_xp': 0,
        'completed_count': 0,
        'total_count': record.total_count,
        'level': _level_payload(0),
        'earned_badge_ids': [],
    })


# =============================================================================
# ADMIN DASHBOARD — an in-app, staff-only view of every user's real activity.
# This is in addition to Django's built-in /admin/ (still fully available);
# this page is styled to match the rest of the app for a supervisor demo.
# =============================================================================
def _is_staff(user):
    return user.is_authenticated and user.is_active and user.is_staff


@login_required
def admin_dashboard(request):
    if not _is_staff(request.user):
        return HttpResponseForbidden("You don't have permission to view this page.")

    user_rows = []
    for u in User.objects.all().order_by('-date_joined'):
        records = AuditRecord.objects.filter(user=u)
        total_audits = records.count()
        total_xp = sum(r.total_xp for r in records)
        avg_score = round(sum(r.overall for r in records) / total_audits) if total_audits else 0
        user_rows.append({
            'user': u,
            'total_audits': total_audits,
            'total_xp': total_xp,
            'avg_score': avg_score,
            'last_audit': records.first(),
        })

    all_records = AuditRecord.objects.all()
    site_stats = {
        'total_users': User.objects.count(),
        'total_audits': all_records.count(),
        'total_xp': sum(r.total_xp for r in all_records),
    }

    return render(request, 'admin_dashboard.html', {
        'user_rows': user_rows,
        'site_stats': site_stats,
    })


@login_required
def admin_user_detail(request, user_id):
    if not _is_staff(request.user):
        return HttpResponseForbidden("You don't have permission to view this page.")

    target_user = get_object_or_404(User, pk=user_id)
    learner_profile, _ = LearnerProfile.objects.get_or_create(user=target_user)

    if request.method == 'POST':
        target_user.first_name = (request.POST.get('first_name') or '').strip()
        target_user.last_name = (request.POST.get('last_name') or '').strip()
        target_user.email = (request.POST.get('email') or '').strip()
        target_user.is_active = bool(request.POST.get('is_active'))
        target_user.save()

        valid_difficulties = dict(LearnerProfile.DIFFICULTY_CHOICES)
        selected_levels = [v for v in request.POST.getlist('difficulty_levels') if v in valid_difficulties]
        learner_profile.difficulty_levels = [k for k in valid_difficulties if k in selected_levels]
        learner_profile.save()

        messages.success(request, f"Updated {target_user.username}'s account.")
        return redirect('admin_user_detail', user_id=target_user.id)

    stats = _user_activity_stats(target_user)
    return render(request, 'admin_user_detail.html', {
        'target_user': target_user,
        'learner_profile': learner_profile,
        'records': stats['records'],
        'total_audits': stats['total_audits'],
        'total_xp': stats['total_xp'],
        'avg_score': stats['avg_score'],
        'total_missions': stats['total_missions'],
        'badges_earned_count': stats['badges_earned_count'],
    })
